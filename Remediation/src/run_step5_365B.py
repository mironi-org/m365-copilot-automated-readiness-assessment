"""
Step 5: Priority-Based Document Segmentation
From: M365_Copilot_Assessment_Insights_Generator.md (Option B)
Input: <input_excel> (passed as parameter)
Output: <output_directory>/High_Priority_Action_Plan.docx
        <output_directory>/Medium_Priority_Action_Plan.docx
        <output_directory>/Low_Priority_Action_Plan.docx
"""

import os
import sys
from datetime import datetime

try:
    import openpyxl
except ImportError:
    os.system(f"{sys.executable} -m pip install openpyxl -q")
    import openpyxl

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

# ─── CONFIG ───────────────────────────────────────────────────────────────────
if len(sys.argv) < 3:
    print("ERROR: Both parameters are required.")
    print("  <input_excel>       - Path to the assessment Excel file (REQUIRED)")
    print("  <output_directory>  - Path to the output directory (REQUIRED)")
    print()
    print("Usage: python run_step5_365B.py <input_excel> <output_directory>")
    sys.exit(1)

EXCEL_PATH = sys.argv[1]
OUTPUT_DIR = sys.argv[2]

if not os.path.isfile(EXCEL_PATH):
    print(f"ERROR: Input file not found: {EXCEL_PATH}")
    sys.exit(1)

if not OUTPUT_DIR.strip():
    print("ERROR: Output directory cannot be empty.")
    sys.exit(1)

# Priority theme colors
PRIORITY_COLORS = {
    "Critical": RGBColor(192, 0, 0),      # Dark red
    "High": RGBColor(255, 102, 0),         # Orange
    "Medium": RGBColor(255, 192, 0),       # Yellow/Gold
    "Low": RGBColor(146, 208, 80),         # Green
}
NAVY_BLUE = RGBColor(0x1F, 0x4E, 0x78)
LIGHT_BLUE = RGBColor(0x5B, 0x9B, 0xD5)
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
GRAY = RGBColor(0x80, 0x80, 0x80)

PRIORITY_EMOJIS = {"Critical": "🔴", "High": "🟠", "Medium": "🟡", "Low": "🟢"}

PRIORITY_TIMELINES = {
    "Critical": "0-7 days — Immediate security risks blocking Copilot deployment",
    "High": "7-30 days — Essential security and identity configurations",
    "Medium": "30-90 days — Optimization and governance improvements",
    "Low": "90+ days — Long-term enhancements and best practices",
}


# ─── RISK ASSESSMENT LOGIC ───────────────────────────────────────────────────
def assess_risk(feature, observation, status):
    f = (feature or "").lower()
    o = (observation or "").lower()
    s = (status or "").lower()

    fn_risk = "Medium"

    if "conditional access" in f or "mfa" in f or "multi-factor" in f:
        impact = "High — identity security gap affects Copilot access controls"
    elif "sharepoint" in f or "onedrive" in f or "content" in f:
        impact = "High — insufficient content limits Copilot knowledge base"
    elif "defender" in f or "endpoint" in f or "xdr" in f:
        impact = "High — unmanaged devices create data protection blind spots"
    elif "sensitivity label" in f or "dlp" in f or "information protection" in f:
        impact = "High — missing data protection policies expose sensitive data via Copilot"
    else:
        impact = "Medium — may affect Copilot readiness posture"

    if "access denied" in s:
        return {"telemetry": "Unknown", "fp_risk": "Medium", "fn_risk": "Medium", "confidence": "Medium", "impact": impact}
    if "warning" in s:
        return {"telemetry": "Telemetry Partial", "fp_risk": "High", "fn_risk": "High", "confidence": "Low", "impact": impact}
    if "critical" in s:
        return {"telemetry": "Derived Posture", "fp_risk": "Medium", "fn_risk": "High", "confidence": "Medium", "impact": impact}
    if "pendingactivation" in s:
        return {"telemetry": "Permission Blocked", "fp_risk": "High", "fn_risk": "Medium", "confidence": "Low", "impact": impact}
    if "pendinginput" in s:
        return {"telemetry": "Manual Verification Required", "fp_risk": "High", "fn_risk": "Medium", "confidence": "Low", "impact": impact}
    if "missing prerequisite" in s or "permission required" in s:
        return {"telemetry": "Unknown", "fp_risk": "Medium", "fn_risk": "Medium", "confidence": "Medium", "impact": impact}
    if "disabled" in s:
        return {"telemetry": "Configuration Gap", "fp_risk": "Medium", "fn_risk": fn_risk, "confidence": "Medium", "impact": impact}
    if "action required" in s:
        return {"telemetry": "Configuration Gap", "fp_risk": "Medium", "fn_risk": fn_risk, "confidence": "Medium", "impact": impact}
    if "insight" in s:
        return {"telemetry": "Index Inferred", "fp_risk": "Medium", "fn_risk": fn_risk, "confidence": "Medium", "impact": impact}
    return {"telemetry": "Index Inferred", "fp_risk": "Medium", "fn_risk": fn_risk, "confidence": "Medium", "impact": impact}


# ─── SUCCESS CRITERIA ─────────────────────────────────────────────────────────
def get_success_criteria(feature):
    f = (feature or "").lower()
    if "conditional access" in f or "ca polic" in f:
        return ["CA policy created and enabled", "Policy targets appropriate user scope", "Required grant controls enforced", "Policy tested in Report-only for 7+ days", "Sign-in logs confirm enforcement", "Break-glass accounts excluded"]
    elif "mfa" in f or "multi-factor" in f or "multifactor" in f:
        return ["100% of targeted users enrolled in MFA", "Each user has primary + backup method", "CA policy enforcing MFA is enabled", "Help desk prepared with support procedures", "Zero extended lockout incidents"]
    elif "sharepoint" in f or "content" in f or "site" in f:
        return ["Sites created with proper structure and permissions", "Content migrated with metadata intact", "100+ documents indexed and searchable", "Copilot can reference migrated content", "Governance policies applied", "User training completed"]
    elif "defender" in f or "xdr" in f or "device" in f or "endpoint" in f:
        return ["50%+ devices onboarded to Defender", "Security policies applied", "AIR enabled", "Alert notifications configured", "SOC team trained"]
    elif "sensitivity label" in f or "dlp" in f or "information protection" in f:
        return ["Label taxonomy defined and published", "Auto-labeling policies configured", "DLP policies enforcing", "Acceptable false positive rate validated", "User training distributed"]
    elif "copilot" in f:
        return ["Service activated and accessible to licensed users", "Feature verified working in target apps", "User adoption communication sent", "Monitoring enabled for usage tracking", "Support team briefed on troubleshooting"]
    elif "teams" in f or "meeting" in f or "collaboration" in f:
        return ["Service/feature enabled in Teams admin center", "User access verified", "Policies configured per requirements", "Communication sent to affected users", "Usage monitoring active"]
    elif "entra" in f or "identity" in f or "access" in f:
        return ["Configuration completed in Entra admin center", "Policy targeting verified", "Test/pilot validated without issues", "Documentation updated", "Next assessment expected to pass"]
    elif "intune" in f or "device management" in f:
        return ["Enrollment configured", "Compliance policies assigned", "CA linked to device compliance", "Non-compliant devices tracked", "Reporting established"]
    elif "phone" in f or "audio" in f or "voice" in f:
        return ["Service activated in Teams admin center", "Licenses assigned to target users", "Calling/conferencing features verified", "User communication completed", "Support documentation updated"]
    else:
        return ["Configuration completed as per recommendation", "Changes validated in admin portal", "No adverse user impact observed", "Documentation updated", "Next assessment expected to show resolved status"]


# ─── DOCUMENT HELPERS ─────────────────────────────────────────────────────────
def set_paragraph_format(paragraph, space_before=None, space_after=None, alignment=None):
    if space_before is not None:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        paragraph.alignment = alignment


def add_heading_colored(doc, text, level, color):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading


def add_table_row(table, label, value, bold_label=True):
    row = table.add_row()
    cell_label = row.cells[0]
    cell_value = row.cells[1]
    p_label = cell_label.paragraphs[0]
    run = p_label.add_run(label)
    if bold_label:
        run.bold = True
    cell_value.paragraphs[0].add_run(str(value) if value else "")


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    c = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
    u = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = paragraph._element.makeelement(qn('w:t'), {})
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_footer(doc, text):
    """Add footer text to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─── GENERATE A SINGLE PRIORITY DOCUMENT ─────────────────────────────────────
def generate_priority_doc(priority_level, items, output_dir):
    """Generate a self-contained Word document for one priority level."""
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    color = PRIORITY_COLORS.get(priority_level, RGBColor(0, 0, 0))
    emoji = PRIORITY_EMOJIS.get(priority_level, "⚪")
    timeline = PRIORITY_TIMELINES.get(priority_level, "As capacity allows")
    now = datetime.now().strftime("%B %d, %Y")

    # ─── 1. COVER PAGE ───────────────────────────────────────────────────────
    # Title
    title_p = doc.add_paragraph()
    set_paragraph_format(title_p, space_before=72, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = title_p.add_run(f"{emoji} {priority_level} Priority Items")
    run.font.size = Pt(28)
    run.font.color.rgb = color
    run.bold = True

    # Subtitle
    sub1 = doc.add_paragraph()
    set_paragraph_format(sub1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = sub1.add_run("Microsoft 365 Copilot Readiness")
    run.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    set_paragraph_format(sub2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = sub2.add_run("Detailed Action Plan")
    run.font.size = Pt(16)

    # Metadata
    doc.add_paragraph()
    meta1 = doc.add_paragraph()
    set_paragraph_format(meta1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    meta1.add_run(f"Generated: {now}")

    meta2 = doc.add_paragraph()
    set_paragraph_format(meta2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = meta2.add_run(f"Total Items: {len(items)}")
    run.bold = True
    run.font.size = Pt(14)

    # Timeline
    doc.add_paragraph()
    tl = doc.add_paragraph()
    set_paragraph_format(tl, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = tl.add_run(f"⏱️ {timeline}")
    run.italic = True

    doc.add_page_break()

    # ─── 2. OVERVIEW SECTION ─────────────────────────────────────────────────
    h1 = add_heading_colored(doc, "Overview of Items", 1, NAVY_BLUE)

    desc = doc.add_paragraph()
    desc.add_run(
        f"This document contains {len(items)} {priority_level} priority items identified during the "
        f"Microsoft 365 Copilot readiness assessment. Each item includes the observation, "
        f"recommendation, risk assessment, implementation steps, and success criteria."
    )

    # Summary table
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Item #", "Service", "Feature", "Status"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, item in enumerate(items, 1):
        row = table.add_row()
        row.cells[0].paragraphs[0].text = str(idx)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].text = item["service"]
        feature_text = item["feature"][:50] + "..." if len(item["feature"]) > 50 else item["feature"]
        row.cells[2].paragraphs[0].text = feature_text
        row.cells[3].paragraphs[0].text = item["status"]

    doc.add_page_break()

    # ─── 3. DETAILED ACTION ITEMS ────────────────────────────────────────────
    add_heading_colored(doc, "Detailed Action Items", 1, NAVY_BLUE)

    for idx, item in enumerate(items, 1):
        feature = item["feature"]
        service = item["service"]
        status = item["status"]
        observation = item["observation"]
        recommendation = item["recommendation"]
        link_text = item["link_text"]
        link_url = item["link_url"]

        risk = assess_risk(feature, observation, status)

        # A. Item Header
        h2 = add_heading_colored(doc, f"Item {idx} — {feature}", 2, color)

        # B. Information Table
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = "Table Grid"
        add_table_row(info_table, "Service Category", service)
        add_table_row(info_table, "Priority", priority_level)
        add_table_row(info_table, "Status", status)
        add_table_row(info_table, "False Positive Risk", risk["fp_risk"])
        add_table_row(info_table, "False Negative Risk", risk["fn_risk"])
        add_table_row(info_table, "Confidence Level", risk["confidence"])

        doc.add_paragraph()

        # C. Current Situation
        h3 = add_heading_colored(doc, "📋 Current Situation", 3, LIGHT_BLUE)
        sit_p = doc.add_paragraph()
        sit_p.paragraph_format.left_indent = Inches(0.25)
        sit_p.add_run(observation if observation else "No observation recorded.")

        # D. Recommended Action
        h3 = add_heading_colored(doc, "✅ Recommended Action", 3, LIGHT_BLUE)
        rec_p = doc.add_paragraph()
        rec_p.paragraph_format.left_indent = Inches(0.25)
        rec_p.add_run(recommendation if recommendation else "See Microsoft Learn documentation for guidance.")

        # E. Risk Assessment
        h3 = add_heading_colored(doc, "⚠️ Risk Assessment", 3, LIGHT_BLUE)
        bullets = [
            f"Telemetry Limitation: {risk['telemetry']}",
            f"False Positive Risk: {risk['fp_risk']}",
            f"False Negative Risk: {risk['fn_risk']}",
            f"Confidence Level: {risk['confidence']}",
        ]
        for b in bullets:
            bp = doc.add_paragraph(b, style="List Bullet")

        # F. Reference Documentation
        h3 = add_heading_colored(doc, "📚 Reference Documentation", 3, LIGHT_BLUE)
        ref_p = doc.add_paragraph()
        ref_p.paragraph_format.left_indent = Inches(0.25)
        if link_url:
            ref_p.add_run(f"{link_text or 'Documentation'}: " if link_text else "Reference: ")
            add_hyperlink(ref_p, link_url, link_url)
        else:
            ref_p.add_run("See Microsoft Learn documentation for detailed guidance.")

        # G. Success Criteria
        h3 = add_heading_colored(doc, "✓ Success Criteria", 3, LIGHT_BLUE)
        criteria = get_success_criteria(feature)
        for c in criteria:
            cp = doc.add_paragraph(f"☐ {c}")
            cp.paragraph_format.left_indent = Inches(0.25)

        # H. Implementation Notes
        h3 = add_heading_colored(doc, "📝 Implementation Notes", 3, LIGHT_BLUE)
        notes_table = doc.add_table(rows=0, cols=2)
        notes_table.style = "Table Grid"
        add_table_row(notes_table, "Assigned To", "")
        add_table_row(notes_table, "Target Completion Date", "")
        add_table_row(notes_table, "Status", "☐ Not Started  ☐ In Progress  ☐ Completed  ☐ Blocked")
        add_table_row(notes_table, "Notes", "")

        # I. Item Separator
        if idx < len(items):
            doc.add_paragraph()
            sep = doc.add_paragraph()
            sep.add_run("_" * 80)
            doc.add_paragraph()

    # ─── 4. APPENDIX ─────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading_colored(doc, "Appendix", 1, NAVY_BLUE)

    # Section A — Related Resources
    add_heading_colored(doc, "Related Resources", 2, color)
    resources = [
        ("Microsoft 365 Copilot Documentation", "https://learn.microsoft.com/microsoft-365-copilot/"),
        ("Entra Admin Center", "https://entra.microsoft.com"),
        ("Microsoft 365 Admin Center", "https://admin.microsoft.com"),
        ("Security Portal", "https://security.microsoft.com"),
        ("Teams Admin Center", "https://admin.teams.microsoft.com"),
        ("SharePoint Admin Center", "https://admin.microsoft.com/sharepoint"),
    ]
    for name, url in resources:
        rp = doc.add_paragraph()
        rp.add_run(f"{name}: ")
        add_hyperlink(rp, url, url)

    # Section B — Priority Timeline Reference
    doc.add_paragraph()
    add_heading_colored(doc, "Priority Timeline Reference", 2, color)
    for p, tl in PRIORITY_TIMELINES.items():
        e = PRIORITY_EMOJIS[p]
        tp = doc.add_paragraph(f"{e} {p}: {tl}", style="List Bullet")

    # ─── FOOTER ──────────────────────────────────────────────────────────────
    add_footer(doc, "Microsoft 365 Copilot Readiness - Action Plan")

    # ─── SAVE ────────────────────────────────────────────────────────────────
    filename = f"{priority_level}_Priority_Action_Plan.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    size = os.path.getsize(filepath)
    print(f"  ✓ {filename} — {len(items)} items, {size:,} bytes")
    return filepath


# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Load Excel
    print("Loading assessment data...")
    wb = openpyxl.load_workbook(EXCEL_PATH, read_only=True, data_only=True)
    ws = wb.active

    headers = [str(cell.value or "").strip() for cell in ws[1]]

    def find_col(keywords):
        for h in headers:
            for kw in keywords:
                if kw.lower() in h.lower():
                    return h
        return None

    col_service = find_col(["service"]) or "Service"
    col_feature = find_col(["feature"]) or "Feature"
    col_status = find_col(["status"]) or "Status"
    col_priority = find_col(["priority"]) or "Priority"
    col_observation = find_col(["observation"]) or "Observation"
    col_recommendation = find_col(["recommendation"]) or "Recommendation"
    col_link_text = find_col(["link text", "linktext"]) or "Link Text"
    col_link_url = find_col(["link url", "linkurl"]) or "Link URL"

    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row and any(cell is not None for cell in row):
            row_dict = {}
            for i, h in enumerate(headers):
                val = row[i] if i < len(row) else None
                row_dict[h] = str(val).strip() if val is not None else ""
            rows.append(row_dict)
    wb.close()

    # Filter non-success
    action_items = [r for r in rows if r.get(col_status, "").lower() and "success" not in r.get(col_status, "").lower()]

    # Group by priority
    priority_groups = {}
    for item in action_items:
        p = item.get(col_priority, "") or "Unknown"
        if p not in priority_groups:
            priority_groups[p] = []
        priority_groups[p].append({
            "service": item.get(col_service, ""),
            "feature": item.get(col_feature, ""),
            "status": item.get(col_status, ""),
            "observation": item.get(col_observation, ""),
            "recommendation": item.get(col_recommendation, ""),
            "link_text": item.get(col_link_text, ""),
            "link_url": item.get(col_link_url, ""),
        })

    print(f"\nTotal action items: {len(action_items)}")
    print(f"Priority distribution:")
    for p in ["Critical", "High", "Medium", "Low"]:
        if p in priority_groups:
            print(f"  {PRIORITY_EMOJIS.get(p, '⚪')} {p}: {len(priority_groups[p])} items")

    # Generate documents
    print(f"\nGenerating priority documents...")
    generated = []
    for priority in ["Critical", "High", "Medium", "Low"]:
        if priority in priority_groups and len(priority_groups[priority]) > 0:
            filepath = generate_priority_doc(priority, priority_groups[priority], OUTPUT_DIR)
            generated.append((priority, len(priority_groups[priority]), filepath))

    # Summary
    print(f"\n{'='*60}")
    print(f"STEP 5 COMPLETE — Priority-Based Document Segmentation")
    print(f"{'='*60}")
    print(f"Documents created: {len(generated)}")
    print(f"Total items across all documents: {len(action_items)}")
    print(f"Output directory: {OUTPUT_DIR}")
    print(f"\nBreakdown:")
    for priority, count, filepath in generated:
        print(f"  {PRIORITY_EMOJIS[priority]} {priority}: {count} items → {os.path.basename(filepath)}")
    print(f"\nFormatting features applied:")
    print(f"  ✓ Priority theme colors (per-document)")
    print(f"  ✓ Cover page with timeline and metadata")
    print(f"  ✓ Summary overview table")
    print(f"  ✓ Per-item: info table, observation, recommendation, risk, criteria, notes")
    print(f"  ✓ Clickable reference links")
    print(f"  ✓ Appendix with resources and timeline reference")
    print(f"  ✓ Footer on all pages")


if __name__ == "__main__":
    main()
