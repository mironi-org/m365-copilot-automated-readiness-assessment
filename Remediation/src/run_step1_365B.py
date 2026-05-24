"""
Step 1: Enterprise Assessment Data Analysis & Executive Summary
From: M365_Copilot_Assessment_Insights_Generator.md (Option B)
Input: <input_excel> (passed as parameter)
Output: <output_directory> (passed as parameter)
  - Executive_Summary.docx
  - status_distribution.png + .json
  - priority_distribution.png + .json
  - top_services.png + .json
  - action_required_by_service.png + .json
"""

import os
import sys
import json
from datetime import datetime
from collections import Counter

# Install dependencies if needed
required = ["openpyxl", "python-docx", "matplotlib"]
for pkg in required:
    try:
        __import__(pkg.replace("-", "_"))
    except ImportError:
        os.system(f"{sys.executable} -m pip install {pkg} -q")

import openpyxl
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches


# === CONFIG ===
if len(sys.argv) < 3:
    print("ERROR: Both parameters are required.")
    print("  <input_excel>       - Path to the assessment Excel file (REQUIRED)")
    print("  <output_directory>  - Path to the output directory (REQUIRED)")
    print()
    print("Usage: python run_step1_365B.py <input_excel> <output_directory>")
    sys.exit(1)

EXCEL_PATH = sys.argv[1]
OUTPUT_DIR = sys.argv[2]

if not os.path.isfile(EXCEL_PATH):
    print(f"ERROR: Input file not found: {EXCEL_PATH}")
    sys.exit(1)

if not OUTPUT_DIR.strip():
    print("ERROR: Output directory cannot be empty.")
    sys.exit(1)

NOW = datetime.now().strftime("%B %d, %Y")

# Professional color palette
COLORS = {
    "Critical": "#C62828",
    "High": "#E65100",
    "Medium": "#F9A825",
    "Low": "#2E7D32",
    "Success": "#1565C0",
    "Action Required": "#C62828",
    "Insight": "#6A1B9A",
    "Warning": "#E65100",
    "PendingActivation": "#F57F17",
    "Access Denied": "#455A64",
}

CHART_COLORS = ["#1565C0", "#C62828", "#E65100", "#F9A825", "#2E7D32", "#6A1B9A", "#455A64", "#00838F", "#4E342E", "#283593"]


def load_data(path):
    """Load Excel and return headers + list of dicts."""
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    headers = [str(c.value or "").strip() for c in ws[1]]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row and any(v is not None for v in row):
            d = {}
            for i, h in enumerate(headers):
                val = row[i] if i < len(row) else None
                d[h] = str(val).strip() if val is not None else ""
            rows.append(d)
    wb.close()
    return headers, rows


def find_col(headers, keywords):
    """Find column name by keyword match."""
    for h in headers:
        for kw in keywords:
            if kw.lower() in h.lower():
                return h
    return keywords[0]


def generate_charts(rows, col_status, col_priority, col_service, output_dir):
    """Generate 4 PNG charts + 4 JSON data files."""
    os.makedirs(output_dir, exist_ok=True)

    # 1. Status Distribution (Pie)
    status_counts = Counter(r[col_status] for r in rows if r[col_status])
    fig, ax = plt.subplots(figsize=(8, 6))
    labels = list(status_counts.keys())
    values = list(status_counts.values())
    colors = [COLORS.get(l, "#90A4AE") for l in labels]
    wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
    for t in autotexts:
        t.set_fontsize(9)
    ax.set_title("Assessment Status Distribution", fontsize=14, fontweight='bold')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, "status_distribution.png"), dpi=150, bbox_inches='tight')
    plt.close()
    with open(os.path.join(output_dir, "status_distribution.json"), "w") as f:
        json.dump(dict(status_counts), f, indent=2)

    # 2. Priority Distribution (Bar) - only non-success
    non_success = [r for r in rows if r[col_status].lower() != "success"]
    priority_counts = Counter(r[col_priority] for r in non_success if r[col_priority])
    order = ["Critical", "High", "Medium", "Low"]
    ordered_priorities = [(p, priority_counts.get(p, 0)) for p in order if p in priority_counts]
    # Add any others
    for p, c in priority_counts.items():
        if p not in order:
            ordered_priorities.append((p, c))

    fig, ax = plt.subplots(figsize=(8, 5))
    plabels = [x[0] for x in ordered_priorities]
    pvalues = [x[1] for x in ordered_priorities]
    pcolors = [COLORS.get(l, "#90A4AE") for l in plabels]
    bars = ax.bar(plabels, pvalues, color=pcolors, edgecolor='white', linewidth=0.5)
    ax.set_title("Priority Level Distribution (Non-Success Items)", fontsize=13, fontweight='bold')
    ax.set_ylabel("Count")
    ax.set_xlabel("Priority")
    for bar, val in zip(bars, pvalues):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3, str(val),
                ha='center', va='bottom', fontsize=10, fontweight='bold')
    ax.set_ylim(0, max(pvalues) * 1.2 if pvalues else 10)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, "priority_distribution.png"), dpi=150, bbox_inches='tight')
    plt.close()
    with open(os.path.join(output_dir, "priority_distribution.json"), "w") as f:
        json.dump(dict(ordered_priorities), f, indent=2)

    # 3. Top 10 Services by Observation Count (Bar)
    service_counts = Counter(r[col_service] for r in rows if r[col_service])
    top10 = service_counts.most_common(10)
    fig, ax = plt.subplots(figsize=(10, 6))
    slabels = [x[0] for x in top10]
    svalues = [x[1] for x in top10]
    bars = ax.barh(slabels[::-1], svalues[::-1], color=CHART_COLORS[:len(top10)], edgecolor='white')
    ax.set_title("Top Services by Observation Count", fontsize=13, fontweight='bold')
    ax.set_xlabel("Number of Observations")
    for bar, val in zip(bars, svalues[::-1]):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height()/2, str(val),
                va='center', fontsize=9)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, "top_services.png"), dpi=150, bbox_inches='tight')
    plt.close()
    with open(os.path.join(output_dir, "top_services.json"), "w") as f:
        json.dump(dict(top10), f, indent=2)

    # 4. Action Required Items by Service (Bar)
    action_items = [r for r in rows if "action" in r[col_status].lower() or "required" in r[col_status].lower()]
    action_by_service = Counter(r[col_service] for r in action_items if r[col_service])
    if not action_by_service:
        # Fallback: use all non-success
        action_by_service = Counter(r[col_service] for r in non_success if r[col_service])
    top_action = action_by_service.most_common(10)

    fig, ax = plt.subplots(figsize=(10, 6))
    alabels = [x[0] for x in top_action]
    avalues = [x[1] for x in top_action]
    bars = ax.barh(alabels[::-1], avalues[::-1], color="#C62828", edgecolor='white', alpha=0.85)
    ax.set_title("Action Required Items by Service", fontsize=13, fontweight='bold')
    ax.set_xlabel("Number of Action Items")
    for bar, val in zip(bars, avalues[::-1]):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, str(val),
                va='center', fontsize=9)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, "action_required_by_service.png"), dpi=150, bbox_inches='tight')
    plt.close()
    with open(os.path.join(output_dir, "action_required_by_service.json"), "w") as f:
        json.dump(dict(top_action), f, indent=2)

    return status_counts, priority_counts, service_counts, non_success


def generate_executive_summary(rows, headers, col_service, col_feature, col_status, col_priority,
                                col_observation, col_recommendation, status_counts, priority_counts,
                                service_counts, non_success, output_dir):
    """Generate Executive_Summary.docx."""
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # === COVER PAGE ===
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Microsoft 365 Copilot Readiness Assessment")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Executive Summary Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Generated: {NOW}\n").font.size = Pt(12)
    info.add_run(f"Total Observations: {len(rows)}\n").font.size = Pt(12)
    info.add_run(f"Action Items: {len(non_success)}\n").font.size = Pt(12)
    services_unique = len(set(r[col_service] for r in rows if r[col_service]))
    info.add_run(f"Services Analyzed: {services_unique}").font.size = Pt(12)

    doc.add_page_break()

    # === EXECUTIVE OVERVIEW ===
    doc.add_heading("Executive Overview", level=1)
    total = len(rows)
    success_count = status_counts.get("Success", 0)
    readiness_pct = round((success_count / total) * 100, 1) if total else 0

    doc.add_paragraph(
        f"This executive summary presents the findings from the Microsoft 365 Copilot Readiness Assessment "
        f"conducted for the organization. The assessment analyzed {total} observations across {services_unique} "
        f"service categories. Of these, {success_count} ({readiness_pct}%) met readiness criteria, while "
        f"{len(non_success)} items require attention before full Copilot deployment readiness is achieved."
    )

    # === KEY STATISTICS TABLE ===
    doc.add_heading("Key Statistics", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Medium Shading 1 Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("Total Observations", str(total)),
        ("Services Analyzed", str(services_unique)),
        ("Success Items", f"{success_count} ({readiness_pct}%)"),
        ("Action Items (Non-Success)", str(len(non_success))),
        ("High Priority Items", str(priority_counts.get("High", 0))),
        ("Medium Priority Items", str(priority_counts.get("Medium", 0))),
    ]
    for i, (label, value) in enumerate(stats):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # === READINESS STATUS BREAKDOWN ===
    doc.add_heading("Readiness Status Breakdown", level=1)
    table = doc.add_table(rows=len(status_counts) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Count"
    table.rows[0].cells[2].text = "Percentage"
    for i, (status, count) in enumerate(sorted(status_counts.items(), key=lambda x: -x[1]), 1):
        table.rows[i].cells[0].text = status
        table.rows[i].cells[1].text = str(count)
        table.rows[i].cells[2].text = f"{round(count/total*100, 1)}%"

    doc.add_paragraph()

    # === PRIORITY ANALYSIS ===
    doc.add_heading("Priority Analysis", level=1)
    doc.add_paragraph(
        f"Among the {len(non_success)} non-success items, the priority distribution indicates the "
        f"urgency and resource allocation needed for remediation:"
    )
    table = doc.add_table(rows=len(priority_counts) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.rows[0].cells[0].text = "Priority"
    table.rows[0].cells[1].text = "Count"
    table.rows[0].cells[2].text = "Target Timeline"
    timelines = {"Critical": "Immediate (1-7 days)", "High": "30 days", "Medium": "90 days", "Low": "120 days"}
    order = ["Critical", "High", "Medium", "Low"]
    row_idx = 1
    for p in order:
        if p in priority_counts:
            table.rows[row_idx].cells[0].text = p
            table.rows[row_idx].cells[1].text = str(priority_counts[p])
            table.rows[row_idx].cells[2].text = timelines.get(p, "TBD")
            row_idx += 1

    doc.add_paragraph()

    # === CRITICAL SECURITY GAPS (Top 5 highest priority) ===
    doc.add_heading("Critical & High Priority Security Gaps", level=1)
    critical_high = [r for r in non_success if r[col_priority] in ("Critical", "High")]
    top5 = critical_high[:5] if critical_high else non_success[:5]

    for i, item in enumerate(top5, 1):
        doc.add_heading(f"{i}. {item[col_feature]}", level=2)
        doc.add_paragraph(f"Service: {item[col_service]} | Priority: {item[col_priority]} | Status: {item[col_status]}")
        if item[col_observation]:
            doc.add_paragraph(f"Finding: {item[col_observation]}")
        if item[col_recommendation]:
            doc.add_paragraph(f"Recommendation: {item[col_recommendation]}")

    # === SERVICE CATEGORY ANALYSIS ===
    doc.add_heading("Service Category Analysis", level=1)
    table = doc.add_table(rows=min(len(service_counts), 10) + 1, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    table.rows[0].cells[0].text = "Service"
    table.rows[0].cells[1].text = "Total Observations"
    table.rows[0].cells[2].text = "Action Items"
    action_by_svc = Counter(r[col_service] for r in non_success if r[col_service])
    for i, (svc, cnt) in enumerate(service_counts.most_common(10), 1):
        if i < len(table.rows):
            table.rows[i].cells[0].text = svc
            table.rows[i].cells[1].text = str(cnt)
            table.rows[i].cells[2].text = str(action_by_svc.get(svc, 0))

    doc.add_paragraph()

    # === STRATEGIC RECOMMENDATIONS ===
    doc.add_heading("Strategic Recommendations", level=1)
    recommendations = [
        "Prioritize High-priority items requiring immediate attention — focus on items blocking Copilot activation.",
        "Address identity and access gaps first — Conditional Access and MFA enrollment are foundational for secure Copilot use.",
        "Establish a SharePoint content strategy to ensure Copilot has adequate indexed content for meaningful responses.",
        "Onboard endpoints to Defender for Endpoint to achieve security visibility required for responsible AI deployment.",
        "Configure sensitivity labels and DLP policies to prevent data leakage through Copilot-generated content.",
        "Implement access reviews for privileged roles to limit Copilot's exposure to sensitive administrative data.",
        "Deploy Teams Premium features to enhance meeting intelligence capabilities with Copilot integration.",
        "Create a phased rollout plan: Phase 1 (security foundations, 30 days), Phase 2 (content readiness, 60 days), Phase 3 (Copilot activation, 90 days).",
        "Establish monitoring and reporting cadence — reassess monthly to track progress against this baseline.",
        "Engage stakeholders early with communication templates to set expectations and drive adoption readiness.",
    ]
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f"{i}. {rec}", style='List Number')

    # === SAVE ===
    doc_path = os.path.join(output_dir, "Executive_Summary.docx")
    doc.save(doc_path)
    return doc_path


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"\n{'='*60}")
    print("  STEP 1: Enterprise Assessment Data Analysis")
    print(f"{'='*60}")
    print(f"  Input: {EXCEL_PATH}")
    print(f"  Output: {OUTPUT_DIR}")
    print(f"  Date: {NOW}")
    print(f"{'='*60}\n")

    # Load
    headers, rows = load_data(EXCEL_PATH)
    print(f"  Loaded {len(rows)} rows with {len(headers)} columns")
    print(f"  Columns: {', '.join(headers)}\n")

    # Identify columns
    col_service = find_col(headers, ["Service"])
    col_feature = find_col(headers, ["Feature"])
    col_status = find_col(headers, ["Status"])
    col_priority = find_col(headers, ["Priority"])
    col_observation = find_col(headers, ["Observation"])
    col_recommendation = find_col(headers, ["Recommendation"])

    # Generate charts
    print("  Generating visualizations...")
    status_counts, priority_counts, service_counts, non_success = generate_charts(
        rows, col_status, col_priority, col_service, OUTPUT_DIR
    )

    # Console statistics
    print(f"\n  --- STATISTICAL ANALYSIS ---")
    print(f"  Total observations: {len(rows)}")
    print(f"  Unique services: {len(set(r[col_service] for r in rows if r[col_service]))}")
    print(f"  Status distribution:")
    for s, c in sorted(status_counts.items(), key=lambda x: -x[1]):
        print(f"    {s}: {c} ({round(c/len(rows)*100,1)}%)")
    print(f"  Priority distribution (non-success):")
    for p in ["Critical", "High", "Medium", "Low"]:
        if p in priority_counts:
            print(f"    {p}: {priority_counts[p]}")
    print(f"  Service breakdown (top 5):")
    for svc, cnt in service_counts.most_common(5):
        print(f"    {svc}: {cnt}")
    print()

    # Data quality
    print("  --- DATA QUALITY ---")
    for h in headers:
        nulls = sum(1 for r in rows if not r.get(h, ""))
        if nulls > 0:
            print(f"    {h}: {nulls} empty values ({round(nulls/len(rows)*100,1)}%)")
    print()

    # Generate Word doc
    print("  Generating Executive_Summary.docx...")
    doc_path = generate_executive_summary(
        rows, headers, col_service, col_feature, col_status, col_priority,
        col_observation, col_recommendation, status_counts, priority_counts,
        service_counts, non_success, OUTPUT_DIR
    )

    # Summary
    print(f"\n{'='*60}")
    print("  STEP 1 COMPLETE — FILES GENERATED:")
    print(f"{'='*60}")
    files = os.listdir(OUTPUT_DIR)
    for f in sorted(files):
        fpath = os.path.join(OUTPUT_DIR, f)
        size = os.path.getsize(fpath)
        print(f"    {f} ({size:,} bytes)")
    print(f"\n  Total files: {len(files)}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
